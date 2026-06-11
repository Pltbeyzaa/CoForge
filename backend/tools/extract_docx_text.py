from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip() + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract DOCX text to stdout or file.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, default=None)
    args = parser.parse_args()

    text = extract_docx_text(args.docx)

    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text, encoding="utf-8")
    else:
        print(text, end="")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

